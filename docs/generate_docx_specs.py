import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ("top", "bottom", "left", "right"):
        if side in kwargs:
            edge = OxmlElement(f'w:{side}')
            for key, val in kwargs[side].items():
                edge.set(qn(f'w:{key}'), str(val))
            tcPr.append(edge)

def add_header_cell(row, i, text):
    cell = row.cells[i]
    cell.text = text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        run = p.runs[0]
    else:
        run = p.add_run()
    run.font.bold = True
    run.font.size = Pt(10)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9D9D9')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Environmental Monitoring Equipment Technical Specifications', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Reference: Circular No. 10/2021/TT-BTNMT dated June 30, 2021 (Vietnam)')
    doc.add_paragraph('_' * 20)

    # --- TABLE 1 ---
    doc.add_heading('Table 1. Automatic Surface Water Monitoring', level=1)
    t1_data = [
        ("1", "Temperature", "°C", "± 5%", "± 3%", "0 ÷ 80°C", "0.1", "≤ 5 s"),
        ("2", "pH", "-", "± 0.1", "± 0.1", "0 ÷ 14", "0.1", "≤ 5 s"),
        ("3", "Total Suspended Solids (TSS)", "mg/L", "± 5%", "± 2%", "0 ÷ 500", "0.1", "≤ 10 s"),
        ("4", "Chemical Oxygen Demand (COD)", "mg/L", "± 5%", "± 5%", "0 ÷ 100", "0.5", "≤ 15 m"),
        ("5", "Dissolved Oxygen (DO)", "mg/L", "± 5%", "± 5%", "0 ÷ 20", "0.1", "≤ 120 s"),
        ("6", "Nitrate (NO₃⁻)", "mg/L", "± 5%", "± 3%", "0 ÷ 50", "0.5", "≤ 10 m"),
        ("7", "Phosphate (PO₄³⁻)", "mg/L", "± 5%", "± 5%", "0 ÷ 2", "-", "≤ 10 m"),
        ("8", "Ammonium (NH₄⁺)", "mg/L", "± 5%", "± 5%", "0 ÷ 5", "0.2", "≤ 30 m"),
        ("9", "Total Phosphorus (TP)", "mg/L", "± 5%", "± 3%", "0 ÷ 2", "0.1", "≤ 30 m"),
        ("10", "Total Nitrogen (TN)", "mg/L", "± 5%", "± 3%", "0 ÷ 20", "0.1", "≤ 30 m"),
        ("11", "Total Organic Carbon (TOC)", "mg/L", "± 5%", "± 2%", "0 ÷ 100", "0.1", "≤ 30 m"),
    ]
    
    table1 = doc.add_table(rows=1, cols=8)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0]
    headers = ["No.", "Parameter", "Unit", "Acc (% rd)", "Acc (% FS)", "Range", "Res", "Resp"]
    for i, h in enumerate(headers):
        add_header_cell(hdr_cells, i, h)
        
    for row_data in t1_data:
        row_cells = table1.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            p = row_cells[i].paragraphs[0]
            if p.runs: p.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # --- TABLE 2 ---
    doc.add_heading('Table 2. Ambient Air Quality Monitoring', level=1)
    table2 = doc.add_table(rows=1, cols=8)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0]
    for i, h in enumerate(headers):
        add_header_cell(hdr_cells, i, h)
    
    t2_raw = [
        ("1", "Temperature", "°C", "± 5%", "± 5%", "0 ÷ 80°C", "0.1", "≤ 120 s"),
        ("2", "Nitrogen Dioxide (NO₂)", "µg/Nm³", "± 5%", "± 5%", "0 ÷ 500", "0.1", "≤ 300 s"),
        ("", "", "ppb", "± 5%", "± 5%", "0 ÷ 250", "0.1", ""),
        ("3", "Carbon Monoxide (CO)", "µg/Nm³", "± 5%", "± 5%", "0 ÷ 100k", "0.1", "≤ 200 s"),
        ("", "", "ppb", "± 5%", "± 5%", "0 ÷ 85k", "0.1", ""),
        ("4", "Sulfur Dioxide (SO₂)", "µg/Nm³", "± 5%", "± 5%", "0 ÷ 1k", "0.1", "≤ 200 s"),
        ("", "", "ppb", "± 5%", "± 5%", "0 ÷ 400", "0.1", ""),
        ("5", "Ozone (O₃)", "µg/Nm³", "± 5%", "± 5%", "0 ÷ 500", "0.1", "≤ 200 s"),
        ("", "", "ppb", "± 5%", "± 5%", "0 ÷ 250", "0.1", ""),
        ("6", "PM10", "µg/Nm³", "± 5%", "± 3%", "0 ÷ 500", "0.1", "≤ 60 s"),
        ("7", "PM2.5", "µg/Nm³", "± 5%", "± 3%", "0 ÷ 150", "0.1", "≤ 60 s"),
    ]
    for row_data in t2_raw:
        row_cells = table2.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            p = row_cells[i].paragraphs[0]
            if p.runs: p.runs[0].font.size = Pt(9)

    doc.add_paragraph("")
    
    # --- TABLE 3 ---
    doc.add_heading('Table 3. Automatic Wastewater Monitoring', level=1)
    t3_data = [
        ("1", "Flow Rate", "m³/h", "± 5%", "± 5%", "-", "≤ 5 m"),
        ("2", "Temperature", "°C", "± 5%", "± 5%", "0.1", "≤ 5 s"),
        ("3", "Color", "Pt-Co", "± 5%", "± 5%", "-", "≤ 5 m"),
        ("4", "pH", "-", "± 0.2", "± 0.2", "0.1", "≤ 5 s"),
        ("5", "TSS", "mg/L", "± 5%", "± 3%", "0.1", "≤ 10 s"),
        ("6", "COD", "mg/L", "± 5%", "± 3%", "0.5", "≤ 15 m"),
        ("7", "Ammonium (NH₄⁺)", "mg/L", "± 5%", "± 5%", "0.2", "≤ 30 m"),
        ("8", "Total Phos (TP)", "mg/L", "± 5%", "± 3%", "0.1", "≤ 30 m"),
        ("9", "Total Nitro (TN)", "mg/L", "± 5%", "± 3%", "0.1", "≤ 30 m"),
        ("10", "TOC", "mg/L", "± 5%", "± 5%", "0.1", "≤ 30 m"),
        ("11", "Res Chlorine", "mg/L", "± 5%", "± 2%", "0.1", "≤ 30 m"),
    ]
    table3 = doc.add_table(rows=1, cols=7)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0]
    h3 = ["No.", "Parameter", "Unit", "Acc(rd)", "Acc(FS)", "Res", "Resp"]
    for i, h in enumerate(h3):
        add_header_cell(hdr_cells, i, h)
    for row_data in t3_data:
        row_cells = table3.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            p = row_cells[i].paragraphs[0]
            if p.runs: p.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # --- TABLE 5 ---
    doc.add_heading('Table 5. Automatic Stack Emission Monitoring', level=1)
    t5_data = [
        ("1", "Temperature", "°C", "± 5%", "± 5%", "0.1", "≤ 120 s"),
        ("2", "Pressure", "kPa", "± 5%", "± 5%", "-", "≤ 120 s"),
        ("3", "Nitrogen Monoxide (NO)", "mg/m³", "± 5%", "± 5%", "1", "≤ 200 s"),
        ("", "", "ppm", "± 5%", "± 5%", "1", ""),
        ("4", "Nitrogen Dioxide (NO₂)", "mg/m³", "± 5%", "± 5%", "1", "≤ 300 s"),
        ("", "", "ppm", "± 5%", "± 5%", "1", ""),
        ("5", "Carbon Monoxide (CO)", "mg/m³", "± 5%", "± 5%", "1", "≤ 200 s"),
        ("", "", "ppm", "± 5%", "± 5%", "1", ""),
        ("6", "Sulfur Dioxide (SO₂)", "mg/m³", "± 5%", "± 5%", "1", "≤ 200 s"),
        ("", "", "ppm", "± 5%", "± 5%", "1", ""),
        ("7", "Oxygen (O₂)", "%V", "± 0.5%", "± 0.5%", "0.1", "≤ 200 s"),
        ("8", "Hydrogen Sulfide (H₂S)", "mg/m³", "± 5%", "± 5%", "0.1", "≤ 300 s"),
        ("", "", "ppm", "± 5%", "± 5%", "0.1", ""),
        ("9", "Ammonia (NH₃)", "mg/m³", "± 5%", "± 5%", "0.1", "≤ 300 s"),
        ("", "", "ppm", "± 5%", "± 5%", "0.1", ""),
        ("10", "Mercury Vapor (Hg)", "mg/m³", "± 5%", "± 5%", "0.1", "≤ 900 s"),
        ("11", "Particulate Matter (PM)", "mg/m³", "± 10%", "± 5%", "0.1", "≤ 60 s"),
    ]
    table5 = doc.add_table(rows=1, cols=7)
    table5.style = 'Table Grid'
    hdr_cells = table5.rows[0]
    for i, h in enumerate(h3):
        add_header_cell(hdr_cells, i, h)
    for row_data in t5_data:
        row_cells = table5.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            p = row_cells[i].paragraphs[0]
            if p.runs: p.runs[0].font.size = Pt(9)

    output = r'c:\Users\Admin\OneDrive\Project\AI\business-planner\docs\table34_monitoring_parameters_EN.docx'
    doc.save(output)
    print(f"File saved successfully to {output}")

if __name__ == "__main__":
    create_doc()
